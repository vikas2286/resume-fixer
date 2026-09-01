"""Parse resumes (PDF/DOCX) into structured JSON.

Extracts: raw text, line/block geometry, fonts, sizes, section headers,
layout signals (columns, tables, header/footer, scanned-image detection),
plus a best-effort structured resume (name / contacts / sections / bullets).
"""
from __future__ import annotations

import os
import difflib
import re
import sys

# ---------------------------------------------------------------- constants

# Canonical section key -> synonyms (common labels and frequent misspellings).
# Matching is fuzzy / case-insensitive so a header is recognized even when
# phrased differently or slightly misspelled.
SECTION_SYNONYMS = {
    "summary": [
        "summary", "professional summary", "executive summary", "summery",
        "profile", "professional profile", "about", "about me", "about myself",
        "career objective", "objective", "overview", "introduction",
        "professional overview", "highlights", "qualifications summary",
    ],
    "experience": [
        "experience", "work experience", "professional experience",
        "employment", "employment history", "work history", "career history",
        "relevant experience", "professional background", "work",
    ],
    "education": [
        "education", "education and training", "educational background",
        "academic background", "academics", "educational qualifications",
        "education qualification", "academic qualifications", "eduaction",
        "school",
    ],
    "skills": [
        "skills", "skill set", "technical skills", "core skills",
        "key skills", "competencies", "technical competencies", "technologies",
        "tools", "areas of expertise", "expertise", "technical expertise",
        "technical stack", "tech stack", "stack", "proficiencies", "skill summary",
    ],
    "projects": [
        "projects", "project", "personal projects", "academic projects",
        "key projects", "major projects", "project experience", "portfolio",
        "projects developed", "project work",
    ],
    "certifications": [
        "certifications", "certification", "certificates", "licenses",
        "licence", "professional certifications", "courses", "training",
        "credentials", "certifications and courses",
    ],
    "awards": [
        "awards", "award", "achievements", "accomplishments", "honors",
        "honours", "awards and honors", "awards and achievements",
        "achievements and awards", "recognitions",
    ],
    # Kanish: LaTeX small-caps header (cmcsc10) - a real TOP-LEVEL section,
    # never folded into the section that happens to precede it.
    "roles": [
        "roles of responsibility", "role of responsibility",
        "positions of responsibility", "position of responsibility",
        "leadership roles", "responsibilities",
    ],
    # Sarthak: distinct section separate from Relevant Coursework.
    "achievements": [
        "achievements and certifications", "achievements & certifications",
        "achievements and certification",
    ],
}

# Bullet markers commonly used in PDFs: • ‣ ▪ ◦ → but not a lone digit ("1.")
_BULLET_MARKERS = "".join([
    "\u2022", "\u2023", "\u00b7", "\u25aa", "\u25cf", "\u25d8", "\u25d9",
    "\u25e6", "\u2043", "\u2219", "-", "\u2010", "\u2011", "\u2012",
    "\u2013", "\u2014", "\u2015", "*", "\u25cb", "\u25a0", "\u25a1",
    "\u25c6", "\uf0b7", "\uf0d8",
])
_BULLET_RE = re.compile(r"^[" + re.escape(_BULLET_MARKERS) + r"](\s*)(?=\S)")

# Chars stripped from the FRONT of a line when it is consumed as bullet
# CONTENT - every recognized level-1/level-2 marker plus Canva private-use
# variants.  The template renders its own marker glyph, so leaving any of
# these on the text would render doubled bullets.
_MARKER_CHARS = ("\u2022*- \u00b7\u25e6\u25aa\u25ab\u2023\u25cb\u2218"
                 "\u25cf\u2043\u2219\u25d8\u25d9\uf0b7\uf0d8")


def strip_bullet_markers(text: str) -> str:
    """Content view of a bullet line: leading markers removed."""
    return (text or "").lstrip(_MARKER_CHARS).strip()


def _normalize_header(text: str) -> str:
    t = (text or "").strip().lower()
    t = re.sub(r"[:.\u00b7\u2022\-_/\\|]+", " ", t)  # punctuation -> space
    t = re.sub(r"\s+", " ", t).strip()
    return t


def _squeezed(text: str) -> str:
    """text with all whitespace removed - matches letter-spaced headers."""
    return re.sub(r"\s+", "", text)


def match_section_header_strong(text):
    """Header match using ONLY high-confidence signals (exact synonym,
    letter-spaced squeeze, starts-with).  Used by the line-merge pass so a
    wrapped fragment like 'ifications' is never mistaken for a header."""
    raw = (text or "").strip()
    if not raw or len(raw) > 60:
        return None
    if _is_bullet(raw):
        return None
    norm = _normalize_header(raw)
    squeezed = _squeezed(norm)
    for key, syns in SECTION_SYNONYMS.items():
        for syn in syns:
            if norm == syn.strip().lower():
                return key
    for key, syns in SECTION_SYNONYMS.items():
        for syn in syns:
            sq = _squeezed(syn)
            if len(sq) >= 5 and squeezed == sq:
                return key
    if len(norm) <= 30 and not norm.endswith("."):
        for key, syns in SECTION_SYNONYMS.items():
            for syn in syns:
                if norm.startswith(syn.strip().lower() + " "):
                    return key
    return None


def match_section_header_debug(text):
    """Same decision as match_section_header() but additionally returns a
    human-readable reason string explaining WHICH rule fired (or why the
    line was rejected).  Used only by the assignment debug log."""
    raw = (text or "").strip()
    if not raw or len(raw) > 60:
        return None, None, "reject: empty or len>60 (%d)" % len(raw)
    # Leading bullet marker => it is a list item, not a section header.
    if re.match(r"^[\u2022\u2023\u00b7\u25aa\u25cf\u25d8\u25d9\u25e6\u2043\u2219"
                + r"\-\u2010\u2011\u2012\u2013\u2014\u2015*\u25cb\u25a0\u25a1\uf0b7\uf0d8]",
                raw):
        return None, None, "reject: starts with bullet marker"
    # Numbered entry title ("1) AlgoZen - ...", "2. Foo") => list item.
    if re.match(r"^\d{1,2}\s*[\)\.\]]", raw):
        return None, None, "reject: numbered entry title"
    norm = _normalize_header(raw)
    squeezed = _squeezed(norm)

    def _fmt(s):
        return "%r(norm=%r)" % (raw, norm)
    # 1) exact synonym match
    for key, syns in SECTION_SYNONYMS.items():
        for syn in syns:
            if norm == syn.strip().lower():
                return key, raw, "RULE1 exact synonym == '%s' [%s]" % (syn, _fmt(raw))
    # 2) squeezed equality -> letter-spaced "A B O U T M E" ~ "about me"
    for key, syns in SECTION_SYNONYMS.items():
        for syn in syns:
            sq = _squeezed(syn)
            if len(sq) >= 5 and squeezed == sq:
                return key, raw, ("RULE2 squeezed == '%s' (letter-spaced?) [%s]"
                                  % (syn, _fmt(raw)))
    # 3) multi-word header starting with a synonym ("Work Experience ...")
    if len(norm) <= 30 and not norm.endswith("."):
        for key, syns in SECTION_SYNONYMS.items():
            for syn in syns:
                s = syn.strip().lower()
                if norm.startswith(s + " "):
                    return key, raw, ("RULE3 starts-with '%s ' [%s]"
                                      % (s, _fmt(raw)))
    # 4) fuzzy / typo tolerance across all synonyms
    best_key, best_ratio, best_syn = None, 0.0, ""
    for key, syns in SECTION_SYNONYMS.items():
        for syn in syns:
            ratio = difflib.SequenceMatcher(None, norm, syn).ratio()
            if ratio > best_ratio:
                best_ratio, best_key, best_syn = ratio, key, syn
    # The length guard blocks wrap-fragments ('ifications' ~ 'certifications',
    # '\x80 Portfolio' ~ 'portfolio') from masquerading as section headers.
    if best_ratio >= 0.72 and abs(len(norm) - len(best_syn)) <= 3:
        return best_key, raw, ("RULE4 fuzzy %.2f ~ '%s' (len %d vs %d) [%s]"
                               % (best_ratio, best_syn, len(norm),
                                  len(best_syn), _fmt(raw)))
    return None, None, ("reject: best fuzzy %.2f < 0.72 (~'%s') or len-guard "
                        "|%d-%d|>3" % (best_ratio, best_syn, len(norm),
                                       len(best_syn)))


def match_section_header(text):
    """Return (canonical_key | None, original_header_text).

    Tolerates case, punctuation, letter-spacing, misspellings and synonyms;
    returns None for ordinary content lines (never splits a section).
    """
    key, title, _reason = match_section_header_debug(text)
    return key, title


def looks_like_section_header(text: str):
    key, _ = match_section_header(text)
    return key


def _is_bullet(text: str) -> bool:
    return bool(_BULLET_RE.match((text or "").strip()))


def _indented_list_candidate(line) -> bool:
    """True when a line sits in the bullet column with no bullet marker."""
    x0 = line.get("x0_pct")
    if x0 is None or not 0.04 <= x0 <= 0.5:
        return False
    text = (line.get("text") or "").strip()
    if not text or len(text) > 220:
        return False
    if line.get("bold"):
        return False  # bold = likely a title/header, not a bullet
    if looks_like_section_header(text):
        return False
    if EMAIL_RE.search(text) or PHONE_RE.search(text) or LINKEDIN_RE.search(text):
        return False
    return True


def _is_bullet_line(line) -> bool:
    """bullet check on a full line dict: marker OR indentation-based."""
    if _is_bullet(line.get("text", "")):
        return True
    return _indented_list_candidate(line)

CORE_SECTIONS = ("experience", "education", "skills")

EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
PHONE_RE = re.compile(r"(\+?\d[\d\s().-]{7,}\d)")
LINKEDIN_RE = re.compile(r"(linkedin\.com/\S+|github\.com/\S+)", re.IGNORECASE)
# Bare hyperlink-label bullets ("Link", "Demo:", "Live") - the target URL
# lives in a PDF link annotation, never as visible text.
_LINK_LABEL_RE = re.compile(
    r"\[?(?:link|demo|live|deployed|deployment|url|(?:view\s+)?project"
    r"|report|paper|code|website|app|github|gitlab|bitbucket)\]?\s*[:\-]?",
    re.IGNORECASE)
# Design-tool exports prepend icon font glyphs (private-use U+E000-F8FF,
# non-printable controls, %-like POI marks) to an otherwise clean chip label
# ("\x1c%� | \x12 Demo" is really a "Demo" chip carrying a hyperlink, but the
# glyph junk made it fail the plain label match and get promoted to an entry
# TITLE - Sarthak's ',% | \x12 Demo' project titles).  Strip that junk before
# matching so chips behave like plain "Demo"/"GitHub" labels everywhere.
_ICON_JUNK_RE = re.compile(
    r"[\x00-\x20\x7f-\x9f\ue000-\uf8ff%\u00a2\u00a3\u00a5\u00b7\u2022\u00a0|]")


def _is_link_chip_line(text: str) -> bool:
    """True when a line is only a bare hyperlink-label chip ("Link",
    "Demo:", "GitHub", possibly with a leading icon glyph run)."""
    t = _ICON_JUNK_RE.sub("", text or "").strip(" \u2013\u2014-:()[]")
    return bool(_LINK_LABEL_RE.fullmatch(t))


def _chip_text(text: str) -> str:
    """Junk-tolerant chip probe: strip icon glyphs AND every non-ascii
    mojibake artifact (Canva subset fonts re-encode '§' as '┬º' etc.),
    then collapse whitespace.  '┬º | \\x12 Demo' -> 'Demo'.  A probe line
    whose visible ascii residue is a bare link label IS a chip."""
    t = _ICON_JUNK_RE.sub("", text or "")
    t = re.sub(r"[^A-Za-z ]", " ", t)
    return re.sub(r"\s+", " ", t).strip()
# Exact-match set of every section synonym, used to keep genuine section
# headers ("PROJECT", "TECH STACK") from being mistaken for link chips.
_SECTION_SYN_LOWER = {s.lower()
                      for syns in SECTION_SYNONYMS.values() for s in syns}


def _is_section_synonym_exact(text: str) -> bool:
    """True when the whole line IS a known section synonym (exact, after
    stripping punctuation) - e.g. 'PROJECT', 'TECH STACK'."""
    t = re.sub(r"[^a-z0-9 ]", "", (text or "").strip().lower()).strip()
    return t in _SECTION_SYN_LOWER


URL_TOKEN_RE = re.compile(
    r"(?:https?://|www\.)[^\s|,\u00b7\u2022]+"
    r"|(?:github\.com|gitlab\.com|linkedin\.com|medium\.com|behance\.net"
    r"|dribbble\.com)/[^\s|,\u00b7\u2022]+",
    re.IGNORECASE)

DATE_RE = re.compile(
    r"\b((19|20)\d{2}\s*(-|\u00b7|\u2010|\u2011|\u2012|\u2013|\u2014|\u2212|to)\s*((19|20)\d{2}|present|current|now)|"
    r"(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*(19|20)?\d{2}\s*"
    r"(-|\u00b7|\u2010|\u2011|\u2012|\u2013|\u2014|\u2212|to)\s*"
    r"((jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*(19|20)?\d{2}|present)"
    # numeric day-first ranges: '22/06/2026 - 22/07/2026', '04.06.2025 - present'.
    # Only the FULL range form is recognized (a lone dd/mm/yyyy is too
    # ambiguous vs. phones/IDs), so this cannot collide with PHONE_RE.
    r"|\d{1,2}[./]\d{1,2}[./](19|20)\d{2}\s*"
    r"(-|\u2010|\u2011|\u2012|\u2013|\u2014|\u2212|to)\s*"
    r"(\d{1,2}[./]\d{1,2}[./](19|20)\d{2}|present|current|now))",
    re.IGNORECASE,
)

# Role-title keywords (Muskan fix, general rule): role lines are
# pattern-predictable ('Summer Internship-3', 'SDE Intern', 'Backend
# Developer') while company names are not.  A title-like line matching this
# set is secondary info (meta) for the current entry, never a NEW entry
# boundary when the entry already has a non-role title.  Word boundaries on
# BOTH ends so 'Engineering'/'Interstitial' never match.
_ROLE_KW_RE = re.compile(
    r"\b(intern|internship|trainee|apprentice|developer|engineer|analyst|"
    r"consultant|associate|designer|architect|manager|lead|freelancer|"
    r"volunteer|fellow|resident|researcher)\b",
    re.IGNORECASE,
)

STANDARD_FONTS = {
    "arial", "helvetica", "calibri", "cambria", "times", "times new roman",
    "georgia", "garamond", "verdana", "tahoma", "trebuchet ms", "lato",
    "open sans", "roboto", "source sans", "noto sans", "segoe ui",
    "liberation serif", "liberation sans", "dejavu", "dejavu sans",
    "dejavu serif",
}


# ---------------------------------------------------------------- helpers

def _norm_font(name) -> str:
    return (name or "").split("+")[-1].strip().lower()


# ---------------------------------------------------------------- PDF parsing

def parse_pdf(path: str) -> dict:
    import fitz  # PyMuPDF

    doc = fitz.open(path)
    lines = []
    fonts = set()
    sizes = []
    full_text_parts = []
    has_tables = False
    n_images_large = 0
    n_images_photo = 0
    # Hyperlink annotations: Canva/LaTeX resumes attach GitHub / LinkedIn /
    # Portfolio as clickable rects whose visible text may be just an icon or a
    # short label - the URL itself never appears in the text stream.
    link_uris = set()
    page_link_rects = []
    link_rects = []
    # Design-tool PDFs (Canva etc.) often draw every glyph twice (fill +
    # overlay), producing identical consecutive lines that would otherwise
    # become duplicate bullets. Collapse them here.
    seen_line_keys = set()

    for pno, page in enumerate(doc):
        pw = page.rect.width or 612
        ph = page.rect.height or 792
        pw_glob[0] = pw
        page_link_rects = []  # per-page: link annotation (uri, rect)
        full_text_parts.append(page.get_text("text"))

        try:
            tables = page.find_tables()
            if tables is not None and len(tables.tables) > 0:
                has_tables = True
        except Exception:
            pass

        try:
            for img in page.get_image_info():
                bb = img.get("bbox", [0, 0, 0, 0])
                if len(bb) == 4 and (bb[2] - bb[0]) * (bb[3] - bb[1]) > 0.5 * pw * ph:
                    n_images_large += 1
                # Photo-like image: >= 40pt in BOTH dimensions. Profile
                # photos / charts / graphics land here; tiny icon glyphs
                # (16-24pt) do not. Feeds the ATS "no_photos" check
                # (stress_photo finding: photo was previously unflagged).
                if (len(bb) == 4 and (bb[2] - bb[0]) >= 40
                        and (bb[3] - bb[1]) >= 40):
                    n_images_photo += 1
        except Exception:
            pass

        try:
            for lk in page.get_links():
                uri = (lk or {}).get("uri")
                if uri and uri.strip():
                    link_uris.add(uri.strip())
                    page_link_rects.append((uri.strip(), lk["from"]))
        except Exception:
            pass

        d = page.get_text("dict")
        for block in d.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                if not spans:
                    continue
                text = "".join(s.get("text", "") for s in spans).strip()
                if not text:
                    continue
                # Normalize unicode dashes so date ranges always match.
                text = re.sub(r"[\u2010-\u2015\u2212]", "-", text)
                main_span = max(spans, key=lambda s: s.get("size", 0))
                size = round(float(main_span.get("size", 0)), 1)
                fname = _norm_font(main_span.get("font"))
                bbox = [round(v, 1) for v in line.get("bbox", [0, 0, 0, 0])]
                norm_text = re.sub(r"\s+", " ", text).lower()
                seen_key = (pno, norm_text, int(bbox[1] // 3))
                if seen_key in seen_line_keys:
                    continue  # double-drawn duplicate at ~same position
                seen_line_keys.add(seen_key)
                fonts.add(fname)
                sizes.append(size)
                # Design-tool PDFs (Canva/Google-Fonts exports) ship glyphs
                # under OBFUSCATED subset font names ("SFBX1095", "SFXC1200")
                # that never contain "bold" - the bold signal lives only in
                # the span flags (bit 16 = bold).  Without it, real bold
                # headers ("Relevant Coursework") look plain and get merged
                # into the previous section (Sarthak finding).  Read BOTH.
                _span_bold = any(int(s.get("flags", 0)) & 16 for s in spans)
                lines.append({
                    "text": text,
                    "font": fname,
                    "bold": (("bold" in fname) or ("black" in fname)
                             or fname.startswith(("cmbx", "cmssbx",
                                                  "eb_garamond12bold"))
                             or _span_bold),
                    "size": size,
                    "page": pno,
                    "bbox": bbox,
                    "y_top_pct": round(bbox[1] / ph, 3) if ph else 0,
                    "y_bot_pct": round(bbox[3] / ph, 3) if ph else 0,
                    "x0_pct": round(bbox[0] / pw, 3) if pw else 0,
                    "bullet": _is_bullet(text),
                })

        # Tag the visible lines each hyperlink overlays ("Link" bullets,
        # "GitHub"/"LinkedIn"/LeetCode icon labels).  Several hyperlinks can
        # share ONE visual row (Canva-style icon contact strips), so a line
        # accumulates MULTIPLE uris - never skip already-tagged lines or all
        # rects after the first come out untagged and vanish from contacts.
        # A rect is PERSONAL when it overlaps ANY top-band line of PAGE 1;
        # anything lower belongs to an entry/project.
        import fitz as _f
        for uri, lr in page_link_rects:
            lr_rect = _f.Rect(lr)
            personal = False
            for ln in lines:
                if ln.get("page") != pno:
                    continue
                lb = _f.Rect(ln["bbox"])
                ir = lb & lr_rect
                if not ir.is_empty and ir.get_area() > 0.3 * min(
                        lb.get_area(), lr_rect.get_area()):
                    tagged = ln.setdefault("link_uris", [])
                    if uri not in tagged:
                        tagged.append(uri)
                    if pno == 0 and ln.get("y_top_pct", 1) < 0.15:
                        personal = True
            link_rects.append({"uri": uri, "page": pno,
                               "bbox": [round(v, 1) for v in lr],
                               "personal": personal})

    raw_text = "\n".join(full_text_parts)
    n_pages = len(doc)
    doc.close()

    # LaTeX/Word PDFs hard-wrap sentences across visual lines (often with
    # hyphenation). Reconstruct logical lines BEFORE anything else, or every
    # wrapped tail becomes a phantom entry/bullet. Design-tool PDFs also need
    # their scrambled block order restored to visual reading order first.
    lines = _sort_reading_order(lines)
    # Multi-column safety net: column-merge and wrapped-line joining both
    # assume single-column reading order. On a genuine 2-column layout they
    # fuse sidebar fragments into false titles and cross-column sentences,
    # scrambling the document (stress_twocol finding). We decide on TWO
    # votes: the pre-merge rail-filtered signal AND the post-merge signal.
    # A single-column resume with right-aligned tags (dates, locations,
    # profile chips) can false-positive the noisy pre-merge vote alone -
    # but then merges CLEAN it up and the post-merge vote is False, so the
    # merges still run. Only when BOTH agree ("genuinely 2-column") do we
    # keep the lines literal: incomplete but never garbled.
    merged = _merge_column_fragments(lines)
    merged = _merge_wrapped_lines(merged)
    _pre_lines = [l for l in lines
                  if lx_bad(l) or not _rail_fragment_ok(l["text"].strip())]
    pre_mc = _layout_signals(_pre_lines, n_pages).get("multicolumn")
    post_mc = _layout_signals(merged, n_pages).get("multicolumn")
    # STRONG multi-column: both votes agree AND the right column actually
    # carries substantial text (>= 20% of characters). A thin chip/tag rail
    # ("LeetCode Profile", right-aligned dates) trips the generic detector
    # but parses perfectly with the normal single-column pipeline, so it
    # must NOT trigger the literal fallback (regression guard: Harsh Arya).
    strong_mc = bool(pre_mc and post_mc and _sidebar_share(lines) >= 0.20)
    if not strong_mc:
        lines = merged
    # Date extraction inside design-tool PDFs can leave bare "()" shells in
    # the line text ("Engineer (2021 - Present)" -> "Engineer ()"). Strip
    # them once, at the line level, so no title/meta/bullet ever renders an
    # empty parenthesis pair (stress finding: "Principal Engineer ()").
    for l in lines:
        t = l["text"]
        if "(" in t and _EMPTY_PARENS_RE.search(t):
            l["text"] = _strip_empty_parens(t)

    parsed = {
        "file_type": "pdf",
        "n_pages": n_pages,
        "raw_text": raw_text,
        "total_chars": len(raw_text.strip()),
        "lines": lines,
        "fonts": sorted(f for f in fonts if f),
        "sizes": sorted(set(sizes)),
        "has_tables": has_tables,
        "large_images": n_images_large,
        "photo_like_images": n_images_photo,
        "links": sorted(link_uris),
        "link_rects": link_rects,
    }
    parsed.update(_layout_signals(lines, n_pages))
    parsed["multicolumn_strong"] = strong_mc
    parsed["contact"] = _detect_contact(parsed)
    parsed["sections"], parsed["section_titles"], parsed["section_order"] = \
        _detect_sections(lines)
    parsed["structured"] = build_structured(parsed)
    return parsed


def _sidebar_share(lines: list) -> float:
    """Fraction of all text characters sitting in the right half of the
    page (x0 > 52%). A genuine sidebar carries a large share (>= 0.2);
    right-aligned chips/dates are a few percent."""
    tot = sum(len(l["text"]) for l in lines)
    if tot <= 0:
        return 0.0
    return (sum(len(l["text"]) for l in lines
                if l.get("x0_pct", 0) > 0.52) / tot)


def _layout_signals(lines: list, n_pages: int) -> dict:
    """Heuristic layout analysis used by both parsers."""
    n = len(lines)
    if n == 0:
        return {"multicolumn": False, "contact_in_header_footer": False,
                "is_scanned": False, "header_lines": [], "footer_lines": [],
                "avg_chars_per_page": 0, "total_text_chars": 0}

    right_col = [l for l in lines if l.get("x0_pct", 0) > 0.52]
    left_col = [l for l in lines if l.get("x0_pct", 0) < 0.38]
    # A genuine side column carries substantial content (a skills list, an
    # education block).  Right-aligned TAG columns - dates, locations,
    # "On-Site"/"Remote" markers, per-entry link labels - are the standard
    # single-column resume pattern (repaired by column-merge), so short
    # fragment lines must not count toward "multicolumn".
    right_content = [l for l in right_col if len(l["text"].split()) >= 3]
    multicolumn = (
        len(right_content) >= 3
        and len(right_col) >= max(4, int(n * 0.12))
        and len(left_col) >= max(4, int(n * 0.12))
    )

    top_lines = [l for l in lines if l.get("y_top_pct", 1) < 0.07]
    header_lines = [l["text"] for l in top_lines if l.get("page", 0) == 0]
    running_headers = [l["text"] for l in top_lines if l.get("page", 0) > 0]
    footer_lines = [l["text"] for l in lines if l.get("y_bot_pct", 0) > 0.93]
    # Contact in a header/footer that ATS tools skip = a real RUNNING header
    # (top band of pages 2+) or a footer.  The page-1 contact block is the
    # standard letterhead, never a "header": flagging it made EVERY multi-page
    # rebuild fail the ATS contact_in_body check, because the old letterhead
    # exemption was skipped whenever page 2 had any line in its top band -
    # which is always true for a 2-page render (Gopal: ATS 95 -> 90 purely
    # from this after "Fix My Resume" produced 2 pages).
    hf_blob = "\n".join(running_headers + footer_lines).lower()
    contact_in_hf = bool(EMAIL_RE.search(hf_blob) or _find_phone(hf_blob))

    return {
        "multicolumn": multicolumn,
        "contact_in_header_footer": contact_in_hf,
        "header_lines": header_lines[:8],
        "footer_lines": footer_lines[:8],
        "avg_chars_per_page": int(sum(len(l["text"]) for l in lines) / max(1, n_pages)),
        "total_text_chars": sum(len(l["text"]) for l in lines),
    }


def _find_phone(text: str):
    """Return a phone-match that isn't actually a date/year-range.

    PHONE_RE happily matches '2017 - 2019' (8 digits, right charset).  Real
    resume phone numbers carry >= 10 digits (+country code), so require that
    and explicitly reject anything DATE_RE claims.
    """
    for m in PHONE_RE.finditer(text):
        tok = m.group(0)
        if len(re.sub(r"\D", "", tok)) >= 10 and not DATE_RE.search(tok):
            return m
    return None


def _detect_contact(parsed: dict) -> dict:
    text = parsed["raw_text"]
    head = "\n".join(text.split("\n")[:12])
    return {
        "email": bool(EMAIL_RE.search(head)),
        "phone": bool(PHONE_RE.search(head)),
        "linkedin": bool(LINKEDIN_RE.search(head)),
    }


def _body_size(lines: list) -> float:
    """Most common font size across a PDF/DOCX - the body text size."""
    if not lines:
        return 11.0
    sizes = {}
    for l in lines:
        s = l.get("size") or 0
        if s:
            sizes[s] = sizes.get(s, 0) + 1
    if not sizes:
        return 11.0
    return max(sizes.items(), key=lambda kv: kv[1])[0]


def _pop_embedded_header(text: str, line: dict):
    """Design-tool PDFs sometimes merge the last content line of one section
    with the NEXT section's header into a single text-stream line (Sarthak:
    '<bullet> Machine Learning' + 'Achievements and Certifications' occupy
    one line).  When such a line ends with a recognized multi-word section
    header phrase, split it: returns (remaining content, header text).

    Guards against false splits in prose bullets ("Won several awards and
    honors"): the line must carry the header's bold signal AND the swallowed
    content fragment must be short (a bullet item, not a sentence).
    """
    raw = (text or "").strip()
    if not _is_bullet(raw) or not line.get("bold"):
        return None
    body = strip_bullet_markers(raw)
    if len(body) < 4:
        return None
    low = body.lower()
    best = None
    for syns in SECTION_SYNONYMS.values():
        for syn in syns:
            # 3+ word phrases only - a trailing lone word ("...awards",
            # "...projects") inside prose must never split.
            if len(syn.split()) < 3 or not low.endswith(syn):
                continue
            cut = len(body) - len(syn)
            head = body[:cut].rstrip()
            header = body[cut:].strip()
            if (head and header
                    and len(head.split()) <= 4 and len(head) <= 45
                    and not head.endswith((":", "-", "\u2013", ","))
                    and match_section_header_strong(header)
                    and (best is None or len(header) > len(best[1]))):
                best = (head, header)
    return best


def _plausible_header_line(line, body_size: float) -> bool:
    """A line that visually reads as a section header even if we can't map it
    to a known category (bold boxed bar, ALL-CAPS bar, elevated font size)."""
    text = (line.get("text") or "").strip()
    if not (2 < len(text) <= 60):
        return False
    if _is_bullet(text):
        return False
    if re.match(r"^\d{1,2}\s*[\)\.\]]", text):
        return False  # numbered entry title ("1) AlgoZen - ...")
    if EMAIL_RE.search(text) or PHONE_RE.search(text) or LINKEDIN_RE.search(text):
        return False
    if DATE_RE.search(text):
        return False
    size = line.get("size") or 0
    bold = bool(line.get("bold"))
    allcaps = text.isupper()
    elevated = size >= body_size + 1.5
    # LaTeX small-caps header font (CM 'cmcsc10') - Kanish finding: headers
    # like 'Roles of Responsibility' are small-caps at body-ish size with NO
    # bold.  A distinct small-caps face is itself a header signal.
    smallcaps = "csc" in (line.get("font") or "").lower()
    if URL_TOKEN_RE.search(text):
        return False  # bare link line ("github.com/user") is contact info
    # A CGPA/percentage score tag is metadata, never a section header
    # (Ashmit finding: the bold right-rail 'CGPA: 9.10' earned bold+allcaps
    # - 'CGPA'.isupper() is True because acronyms uppercase only cased chars
    # - and opened a phantom custom section that carved the education block).
    if _is_metric_tag(text):
        return False
    # A header must be visually distinct from body copy: bold or all-caps or
    # clearly larger. All three at once is noisy; require two signals.
    signals = sum([bold, allcaps, elevated, smallcaps])
    return signals >= 2


def _name_line_text(lines: list):
    """Text of the resume name line: biggest text in the top of page 0."""
    pool = [l for l in lines if l.get("page", 0) == 0][:12] or lines[:12]
    if not pool:
        return None
    return max(pool, key=lambda l: l.get("size", 0))["text"].strip()


def _detect_sections(lines: list) -> tuple:
    """Return (sections, titles, order).

    sections: {section_key: [line, ...]}.
    Known headers map to their canonical key; headers that can't be classified
    are preserved under a ``custom:NN`` key with their ORIGINAL label so their
    content is never silently merged into another section.

    Set RF_DEBUG_ASSIGN=1 to print, for every line, which section it was
    assigned to and why (matched rule / fallback signals / content append).
    """
    dbg = os.environ.get("RF_DEBUG_ASSIGN") == "1"
    body_size = _body_size(lines)
    name_text = _name_line_text(lines)
    if dbg:
        print("\n=== SECTION ASSIGNMENT DEBUG (body_size=%s, name=%r) ==="
              % (body_size, name_text))
    sections = {}
    titles = {}
    order = []
    current = None
    custom_seq = 0

    for l in lines:
        text = l["text"]
        flags = "size=%s bold=%d bullet=%d" % (l.get("size"),
                                               bool(l.get("bold")),
                                               bool(l.get("bullet")))
        if name_text and text.strip() == name_text:
            if dbg:
                print("  SKIP-NAME   %-28s | %s | %r"
                      % ("<name>", flags, text[:60]))
            continue  # the name renders separately, never as a section
        # Hyperlink-label chips ("View Project:", "[Report]:", "Link:")
        # carry URIs - they are ENTRY material, and because several end in
        # words like "Project" they would otherwise classify as a section
        # header and be silently dropped from their section stream.
        # EXCEPTION: a bare word that IS an exact section synonym ("PROJECT",
        # "TECH STACK") is a real header, not a chip - the chip word list
        # contains "project", which otherwise swallows a genuine PROJECT
        # header (stress_fresher finding: whole section leaked into Skills).
        if ((l.get("link_uris") or _LINK_LABEL_RE.fullmatch(text.strip()))
                and not _is_section_synonym_exact(text)):
            if current:
                if dbg:
                    print("  LINK-CHIP   %-28s | %s | -> %s"
                          % ("'%s'" % text.strip()[:26], flags, current))
                sections[current].append(l)
                continue
        # MERGED HEADER SPLIT: a bullet line that swallowed the next
        # section's header ("Machine Learning Achievements and
        # Certifications") contributes its content fragment to the CURRENT
        # section, then the header text continues through normal matching.
        emb = _pop_embedded_header(text, l)
        if emb and current:
            if dbg:
                print("  SPLIT-EMB   %-28s | content=%r header=%r"
                      % ("'%s'" % text.strip()[:26], emb[0][:26], emb[1][:30]))
            sections[current].append(dict(l, text=emb[0], bullet=True))
            l = dict(l, text=emb[1], bullet=False, bold=True)
            text = emb[1]
        key, orig_title = match_section_header(text)
        if key is not None:
            reason = match_section_header_debug(text)[2]
            if dbg:
                print("  HEADER      %-28s | %s | -> %s | %s"
                      % ("'%s'" % orig_title[:26], flags, key, reason))
        else:
            if _plausible_header_line(l, body_size):
                # Unclassified but clearly a header -> keep it as its own
                # labeled section instead of folding content elsewhere.
                key = "custom:%02d" % custom_seq
                custom_seq += 1
                orig_title = text.strip()  # keep the ORIGINAL header label
                size, bold = l.get("size") or 0, bool(l.get("bold"))
                sig = []
                if bold:
                    sig.append("bold")
                if text.isupper():
                    sig.append("allcaps")
                if size >= body_size + 1.5:
                    sig.append("elevated(%.1f>=%.1f+1.5)" % (size, body_size))
                if dbg:
                    print("  FALLBACK    %-28s | %s | -> %s | "
                          "unmapped-but-header: signals=[%s]"
                          % ("'%s'" % orig_title[:26], flags, key,
                             ", ".join(sig)))
            elif dbg:
                why = match_section_header_debug(text)[2]
                print("  CONTENT     %-28s | %s | -> %s | %s"
                      % ("'%s'" % text.strip()[:26], flags,
                         current or "<DROPPED - no open section>", why))
        if key is not None:
            current = key
            if key not in sections:
                sections[key] = []
                titles[key] = orig_title or key
                order.append(key)
            continue
        if current:
            sections[current].append(l)
    if dbg:
        print("--- assignment summary ---")
        for k in order:
            n = len(sections.get(k, []))
            firsts = "; ".join(x["text"][:40] for x in sections.get(k, [])[:3])
            print("  %-14s title=%-24r lines=%d  first: %s"
                  % (k, titles.get(k), n, firsts))
        print("=== END ASSIGNMENT DEBUG ===\n")
    return sections, titles, order


# ------------------------------------------------- wrapped-line reconstruction

_PAGE_NUM_RE = re.compile(r"^\d{1,3}$")
_TERMINATORS = (".", "!", "?", ":")
# Words that leave a sentence syntactically incomplete -> the next short line
# ("Neo4j" after "... knowledge representation in") is a wrap continuation.
_DANGLING_WORDS = {
    "a", "about", "after", "and", "an", "as", "at", "before", "between",
    "by", "for", "from", "in", "into", "like", "of", "on", "onto", "or",
    "over", "such", "than", "the", "through", "to", "under", "using", "via",
    "with", "within", "while", "throughout", "along", "across", "among",
    "during", "including", "towards", "beyond", "upon", "per", "beneath",
}

# Multi-word technical phrases break across visual lines mid-phrase
# ("... Management systems, Object" / "Oriented Programming").  When the
# previous line ENDS with one of these phrase-starting words, the next
# short line is a continuation even if it starts uppercase (C.6/C.7).
_PHRASE_CONTINUERS = _DANGLING_WORDS | {
    "object", "operating", "machine", "data", "database", "web",
    "computer", "software", "artificial", "business", "cloud", "deep",
    "natural", "systems", "science", "learning", "development",
    "management", "analysis", "design", "objectoriented",
}


def _last_word(text: str) -> str:
    parts = text.rstrip().split()
    return parts[-1].lower() if parts else ""


def _sort_reading_order(lines: list) -> list:
    """Put lines into visual reading order (top-to-bottom, left-to-right).

    Design-tool PDFs (Canva / Word exports) emit text blocks in ARBITRARY
    order, which scrambles the sequential section state-machine: an
    'Education' header can appear in the stream AFTER its own content and
    before unrelated project bullets (real bug on 'Harsh Arya Resume').
    Sorting each page by y-band then x fixes the stream while keeping
    float-right dates on the same row as their entry title.
    """
    out = []
    by_page = {}
    for l in lines:
        by_page.setdefault(l.get("page", 0), []).append(l)
    for pno in sorted(by_page):
        out.extend(sorted(
            by_page[pno],
            key=lambda l: (round(l.get("y_top_pct", 0) * 100),
                           l.get("x0_pct", 0)),
        ))
    return out


_METRIC_TAG_RE = re.compile(
    r"(?:\d{1,3}(?:\.\d{1,2})?\s*/\s*(?:10|100)"
    r"|\d{1,3}(?:\.\d{1,2})?\s*%"
    r"|(?:cgpa|gpa|aggregate|percentage|percent)\s*[:\-]?\s*"
    r"\d{1,3}(?:\.\d{1,2})?(?:\s*/\s*(?:10|100))?\s*%?)",
    re.IGNORECASE)


def _is_metric_tag(text: str) -> bool:
    """A standalone CGPA/percentage/score tag ('93%', 'CGPA: 9.07/10').
    Bare numbers are NOT metric tags - too ambiguous to act on."""
    t = (text or "").strip()
    if not t or len(t) > 24 or not any(ch.isdigit() for ch in t):
        return False
    return bool(_METRIC_TAG_RE.fullmatch(t))


def _rail_fragment_ok(text: str) -> bool:
    """Right-rail fragment eligibility: dates, CGPA/percent tags or a SHORT
    single capitalized tag like 'On-Site'.  Long/prose/column-body text,
    anything ending with a sentence terminator and bare LINK-LABEL chips
    ('Link', 'Demo', 'Report' - they carry hyperlinks and are consumed by
    the entry-link rules, never glued into titles) are rejected."""
    t = (text or "").strip()
    if not t or len(t) > 32 or t.endswith(_TERMINATORS):
        return False
    if _is_bullet(t) or match_section_header_strong(t):
        return False
    if _LINK_LABEL_RE.fullmatch(t):
        return False
    if DATE_RE.search(t):
        return True
    # Metric tags: '93%', '93.80%', 'CGPA: 9.07/10', 'CGPA: 8.57', '9.07/10'.
    # The docstring always promised CGPA/percent eligibility but the code
    # never had the branch (Vinod finding: his percent scores sit in the
    # right rail and were rejected, collapsing his 3 education entries into
    # one).  Bare numbers ('10', '530') stay rejected - ambiguous.
    if _is_metric_tag(t):
        return True
    # Single capitalized tag only (e.g. 'On-Site', 'Present') — no spaces,
    # so multi-word prose like 'Hello World' never qualifies as a rail fragment.
    if re.fullmatch(r"[A-Z][A-Za-z0-9&/.'-]{1,15}", t):
        return True
    return False


def _merge_column_fragments(lines: list) -> list:
    """Join same-row left/right fragments produced by design-tool PDFs.

    Canva-style exports emit the degree text and its right-aligned date -
    or a company name and its right-rail tag - as SEPARATE lines even though
    they visually share one row.  Two fragments fuse when they sit on the
    same page with strongly overlapping y-bands, disjoint x-ranges, similar
    font size, and the RIGHT one is a short date/tag token (never prose,
    bullets or headers).  Marker-led rows ("- Freight Tiger | On-Site") are
    deliberately NOT fused here: their title/rail pairing is resolved by
    entry grouping instead, so the tag becomes meta rather than title text.
    The LEFT fragment keeps its style/dict; bboxes union.
    """
    if not lines or len(lines) < 2:
        return lines
    out = []
    used = set()
    pw = pw_glob[0] if pw_glob and pw_glob[0] else 0.0
    for i, l in enumerate(lines):
        if i in used:
            continue
        lb = l.get("bbox") or []
        if (l.get("page") is None or lx_bad(l) or not lb or len(lb) < 4
                or _is_bullet(l["text"]) or match_section_header_strong(
                    l["text"])
                or l["text"].rstrip().endswith(":")):
            out.append(l)
            continue
        lx1 = lb[2] / pw if pw else l.get("x0_pct", 0)
        best_j = None
        for j in range(i + 1, len(lines)):
            if j in used:
                continue
            r = lines[j]
            rb = r.get("bbox") or []
            if (r.get("page") != l.get("page") or lx_bad(r)
                    or not rb or len(rb) < 4):
                continue
            rx0 = r["x0_pct"]
            if rx0 < lx1 + 0.005:          # must start strictly right of us
                continue
            ryov = (min(l["y_bot_pct"], r["y_bot_pct"])
                    - max(l["y_top_pct"], r["y_top_pct"]))
            rh = max(r["y_bot_pct"] - r["y_top_pct"], 1e-4)
            if ryov < 0.55 * rh:
                continue
            if abs((r.get("size") or 0) - (l.get("size") or 0)) > 2:
                continue
            if len(r["text"]) + len(l["text"]) + 1 > 118:
                continue
            if not _rail_fragment_ok(r["text"]):
                continue
            # pure metric tags never fuse: they pair with THEIR OWN row in
            # entry grouping ('BMS Institute...' + 'CGPA: 9.07/10' must stay
            # two lines so the school and the score land in separate fields)
            if _is_metric_tag(r["text"]):
                continue
            best_j = j
            break                          # nearest right neighbour wins
        if best_j is None:
            out.append(l)
            continue
        r = lines[best_j]
        used.add(best_j)
        merged = dict(l)
        merged["bbox"] = [lb[0], lb[1],
                          max(lb[2], rb[2]), max(lb[3], rb[3])]
        ru = r.get("link_uris")
        if ru:
            lu = merged.setdefault("link_uris", [])
            merged["link_uris"] = lu + [u for u in ru if u not in lu]
        merged["text"] = _strip_empty_parens(l["text"].rstrip() + " "
                                             + r["text"].strip())
        merged["col_merged"] = True
        print("[column-merge] %r  <-  %r" % (l["text"][:50],
                                             r["text"][:30]),
              file=sys.stderr)
        out.append(merged)
    return out


def lx_bad(ln: dict) -> bool:
    return ln.get("x0_pct") is None


# page width cache filled by parse_pdf (module-level scratch, single-threaded)
pw_glob = [0.0]


def _union_link_uris(prev: dict, cur: dict) -> None:
    """Carry cur's tagged hyperlink uris into prev across a line merge."""
    cu = cur.get("link_uris")
    if not cu:
        return
    tu = prev.get("link_uris")
    if not tu:
        prev["link_uris"] = list(cu)
        return
    for u in cu:
        if u not in tu:
            tu.append(u)


def _merge_wrapped_lines(lines: list) -> list:
    """Join visual lines that are continuations of the same logical element.

    Merge priority order per line:
      0. strong section header (exact/squeezed/starts-with synonym) or a
         bullet-marker line  -> never merged, starts a new element
      1. lowercase continuation (optionally hyphenated) of the previous line
      2. short bare tail after a dangling word ("... in" + "Neo4j")
    Weak fuzzy header matches do NOT block merging, so a wrapped tail like
    'ifications' rejoins its bullet instead of becoming a phantom header.
    """
    if not lines:
        return lines
    out = [dict(lines[0])]
    for raw in lines[1:]:
        cur = dict(raw)
        if _PAGE_NUM_RE.match(cur["text"]):
            continue  # LaTeX page number
        prev = out[-1]
        same_page = prev.get("page") == cur.get("page")
        cur_has_contact = bool(
            EMAIL_RE.search(cur["text"]) or PHONE_RE.search(cur["text"])
            or LINKEDIN_RE.search(cur["text"]))
        starts_lower = bool(cur["text"]) and cur["text"][:1].islower()
        hyphenated = prev["text"].endswith("-")
        ends_sentence = prev["text"].endswith(_TERMINATORS)
        prev_has_contact = bool(
            EMAIL_RE.search(prev["text"]) or PHONE_RE.search(prev["text"])
            or LINKEDIN_RE.search(prev["text"]))

        # 0. strong header / bullet marker -> never merged
        if match_section_header_strong(cur["text"]) or _is_bullet(cur["text"]):
            out.append(cur)
            continue

        # 1. lowercase continuation (possibly hyphenated)
        if same_page and not prev_has_contact and not cur_has_contact \
                and starts_lower and (hyphenated or not ends_sentence) \
                and len(prev["text"]) + len(cur["text"]) < 240:
            if hyphenated:
                prev["text"] = prev["text"][:-1] + cur["text"]  # dehyphenate
            else:
                prev["text"] = prev["text"] + " " + cur["text"]
            _union_link_uris(prev, cur)
            continue

        # 2. short bare tail after a dangling word ("... in" + "Neo4j")
        #    A lone short word (even one ending in '.', like "Neo4j.") after
        #    a dangling word is a wrapped fragment, not a sentence.
        single_frag = (len(cur["text"]) <= 14
                       and " " not in cur["text"].strip()
                       and re.search(r"[A-Za-z]", cur["text"]))
        cur_standalone = bool(
            EMAIL_RE.search(cur["text"]) or PHONE_RE.search(cur["text"])
            or LINKEDIN_RE.search(cur["text"])
            or re.search(r":@(),\-\u2013\u2014|\b\d+(?:\.\d+)?\b", cur["text"])
            or cur["text"].endswith(_TERMINATORS)) and not single_frag
        if same_page and not prev_has_contact and not cur_has_contact \
                and (not starts_lower
                     or _last_word(prev["text"]) in _PHRASE_CONTINUERS) \
                and not ends_sentence \
                and len(cur["text"]) <= 35 and not cur_standalone \
                and not match_section_header_strong(cur["text"]) \
                and _last_word(prev["text"]) in _PHRASE_CONTINUERS \
                and len(prev["text"]) + len(cur["text"]) < 220:
            prev["text"] = prev["text"] + " " + cur["text"]
            _union_link_uris(prev, cur)
            continue
        out.append(cur)
    for l in out:
        l["bullet"] = _is_bullet(l["text"])
        l["text"] = re.sub(r"\s+", " ", l["text"]).strip()
    return [l for l in out if l["text"]]


# ---------------------------------------------------------------- DOCX parsing

def parse_docx(path: str) -> dict:
    from docx import Document

    document = Document(path)
    lines = []
    fonts = set()
    sizes = []

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        run_sizes, run_fonts = [], []
        for run in para.runs:
            if run.font.size is not None:
                run_sizes.append(run.font.size.pt)
            if run.font.name:
                run_fonts.append(_norm_font(run.font.name))
        size = round(max(run_sizes) if run_sizes else 11.0, 1)
        fname = run_fonts[0] if run_fonts else ""
        fonts.add(fname)
        sizes.append(size)
        lines.append({
            "text": text,
            "font": fname,
            "bold": any(r.bold for r in para.runs) if para.runs else False,
            "size": size,
            "page": 0,
            "bbox": [],
            "y_top_pct": 0.5,
            "y_bot_pct": 0.5,
            # DOCX has no real column geometry; 1.0 keeps indentation-based
            # bullet heuristics (which need real x0 data) from misfiring.
            "x0_pct": 1.0,
            "bullet": _is_bullet(text)
                      or (para.style.name or "").lower().startswith("list"),
        })

    raw_text = "\n".join(l["text"] for l in lines)
    est_pages = max(1, round(len(raw_text) / 3200))
    parsed = {
        "file_type": "docx",
        "n_pages": est_pages,
        "raw_text": raw_text,
        "total_chars": len(raw_text),
        "lines": lines,
        "fonts": sorted(f for f in fonts if f),
        "sizes": sorted(set(sizes)),
        "has_tables": len(document.tables) > 0,
        "large_images": 0,
        "photo_like_images": 0,
        # DOCX is a flow layout: assume single column, body-text contacts.
        "multicolumn": False,
        "contact_in_header_footer": False,
        "header_lines": [],
        "footer_lines": [],
        "avg_chars_per_page": int(len(raw_text) / est_pages),
        "total_text_chars": len(raw_text),
    }
    parsed["sections"], parsed["section_titles"], parsed["section_order"] = \
        _detect_sections(lines)
    parsed["structured"] = build_structured(parsed)
    return parsed


# ---------------------------------------------------------------- public API

def parse_file(path: str) -> dict:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return parse_pdf(path)
    if ext in (".docx", ".doc"):
        return parse_docx(path)
    raise ValueError("Unsupported file type: %s" % ext)


def is_scanned(parsed: dict) -> bool:
    """Image-only PDF with (almost) no selectable text."""
    return parsed["total_chars"] < 60 and parsed.get("large_images", 0) > 0


# ---------------------------------------------------------------- structuring
# Turns flat parsed lines into a clean resume JSON usable by templates.
# If a Gemini key is configured, llm_service.structure_resume() overrides this.


# --------------------------- header-spelling correction (NARROW exception) --
# The ONLY place source-fidelity is overridden: when a DETECTED SECTION
# HEADER fuzzy-matches a canonical section name within a small edit
# distance (typos like 'EDUACTION' / 'Certifcations'), the corrected
# canonical spelling is rendered instead.  Never applied to body content,
# entry titles, company names, coursework, or any other resume text.
_CANONICAL_HEADER_LABELS = (
    "Summary", "Experience", "Education", "Skills", "Projects",
    "Certifications", "Awards", "Achievements",
)


def _damerau_levenshtein(a: str, b: str) -> int:
    """Edit distance with adjacent transpositions counted as 1
    ('eduaction' -> 'education' is one swap = 1, not 2 substitutions)."""
    la, lb = len(a), len(b)
    prev2 = None
    prev = list(range(lb + 1))
    for i in range(1, la + 1):
        cur = [i] + [0] * lb
        for j in range(1, lb + 1):
            cost = 0 if a[i - 1] == b[j - 1] else 1
            cur[j] = min(prev[j] + 1,        # deletion
                         cur[j - 1] + 1,     # insertion
                         prev[j - 1] + cost)  # substitution
            if (i > 1 and j > 1 and a[i - 1] == b[j - 2]
                    and a[i - 2] == b[j - 1]):
                cur[j] = min(cur[j], prev2[j - 2] + 1)  # transposition
        prev2, prev = prev, cur
    return prev[lb]


def canonical_header_spelling(label: str) -> str:
    """Return the corrected canonical spelling when ``label`` is a section
    header within edit distance <= 2 of exactly one canonical name
    ('EDUACTION' -> 'Education').  Correctly-spelled headers, differently
    WORDED headers ('Technical Skills', 'Relevant Coursework') and anything
    longer than a plausible typo return unchanged."""
    norm = re.sub(r"\s+", " ", (label or "")).strip().lower()
    if not norm or len(norm) > 25:
        return label
    best, best_d = None, 99
    for canon in _CANONICAL_HEADER_LABELS:
        c = canon.lower()
        if abs(len(norm) - len(c)) > 2:
            continue
        d = _damerau_levenshtein(norm, c)
        if d < best_d:
            best, best_d = canon, d
    if best and 0 < best_d <= 2:
        return best
    return label


# Section headers some templates write GENERICALLY ('Portfolio', 'Works') for
# what is structurally the projects bucket.  Those render with the canonical
# "Projects" label (the renderer buckets by WHAT a section IS, not how the
# template named it); every OTHER section stays source-faithful.
_PROJECT_GENERIC_ALIASES = {
    "portfolio", "port folio", "my portfolio", "portfolio projects",
    "projects portfolio", "works", "my work", "featured projects",
}


def _sanitize_header_label(label: str) -> str:
    """Strip control/icon prefixes ('\\x80 Portfolio') and stray padding from a
    detected source header so it renders cleanly.  Canva/design-tool PDFs glue
    an icon glyph or Windows-1252 control byte in front of the real label."""
    if not label:
        return label
    # keep C0 + C1 control characters out (U+0080 is a C1 control that
    # Canva exports as a leading icon-glyph byte in front of the label)
    s = "".join(ch for ch in label
                if not (0x00 <= ord(ch) <= 0x1F or 0x7F <= ord(ch) <= 0x9F))
    s = re.sub(r"\s+", " ", s).strip()
    return s.strip(" \u00a0\u2022\u2023\u25e6\u00b7-|\u2013\u2014\u00b0")


def build_structured(parsed: dict) -> dict:
    lines = parsed["lines"]
    if not lines:
        return {"name": "", "headline": "", "contacts": [], "sections": []}

    # --- name: biggest text near the top ----------------------------------
    pool = [l for l in lines if l.get("page", 0) == 0][:12] or lines[:12]
    name_line = max(pool, key=lambda l: l.get("size", 0))
    name = (name_line["text"].strip().title()
            if name_line["text"].isupper() else name_line["text"].strip())

    # --- headline: second-biggest nearby line ------------------------------
    rest = [l for l in pool if l is not name_line]
    headline = ""
    if rest:
        cand = max(rest, key=lambda l: l.get("size", 0))
        if (cand.get("size", 0) >= name_line.get("size", 0) - 3
                and not EMAIL_RE.search(cand["text"])
                and not PHONE_RE.search(cand["text"])
                and not looks_like_section_header(cand["text"])
                and len(cand["text"]) < 80):
            headline = cand["text"].strip()
        else:
            headline = ""

    # --- contacts ----------------------------------------------------------
    contacts = []
    seen_contacts = set()

    def _add_contact(tok):
        tok = tok.strip().rstrip(".,;|")
        if not tok:
            return
        # Wrapped lines can leave newlines INSIDE a URL/contact token
        # ("https://ma\r\narketplace..."). Links have no legal whitespace.
        if URL_TOKEN_RE.search(tok) and not EMAIL_RE.search(tok):
            tok = re.sub(r"\s+", "", tok)
        else:
            tok = re.sub(r"\s+", " ", tok)
        key = re.sub(r"\s+", "", tok).lower()
        if key and key not in seen_contacts:
            seen_contacts.add(key)
            contacts.append(tok)

    for l in pool:
        t = l["text"].strip()
        em = EMAIL_RE.search(t)
        ph = _find_phone(t)
        li = LINKEDIN_RE.search(t)
        if not (em or ph or li):
            # still scan for bare link tokens ("github.com/user", portfolio URLs)
            for m in URL_TOKEN_RE.finditer(t):
                _add_contact(m.group(0))
            continue
        # Pull out only the matched tokens so FontAwesome/icon glyph junk
        # ("# vli7k2a6s@...", "\x83 +91 ...") never leaks into the contact.
        # Email and phone are collected INDEPENDENTLY - '(em or ph or li)'
        # used to shadow the phone whenever it shared a line with the email.
        for m in (em, ph, li):
            if m:
                _add_contact(m.group(0))
        for m in URL_TOKEN_RE.finditer(t):
            _add_contact(m.group(0))

    # Phones hard-wrapped across header lines ("+91 70065190" / "074") are
    # invisible to per-line matching - retry on the whitespace-collapsed blob.
    joined_head = re.sub(r"\s+", " ", " ".join(l["text"] for l in pool))
    jm = _find_phone(joined_head)
    if jm:
        _add_contact(re.sub(r"\s+", " ", jm.group(0)))

    # Hyperlink annotations (GitHub / LinkedIn / Portfolio icons): these URLs
    # exist only as link rects, never as text.  ONLY annotations overlapping
    # the header pool lines are PERSONAL links - project-specific ones
    # ("Link" bullets mid-page) attach to their entry via line tagging and
    # must never leak into the contact header.  mailto:/tel: collapse to
    # address form; design-tool housekeeping links are dropped.
    _JUNK_LINKS = re.compile(
        r"(?:canva\.com|adobe\.com|fonts\.google|create\.pdf)", re.IGNORECASE)
    for lr in parsed.get("link_rects", []):
        if not lr.get("personal"):
            continue  # project/mid-page link -> belongs to its entry, not here
        uri = lr["uri"]
        u = re.sub(r"^(mailto:|tel:)", "", uri, flags=re.IGNORECASE)
        if _JUNK_LINKS.search(u):
            continue
        if EMAIL_RE.fullmatch(u) or re.sub(r"\s+", "", u).isdigit():
            _add_contact(u)
            continue
        if URL_TOKEN_RE.search(u):
            _add_contact(u)

    # --- sections -----------------------------------------------------------
    detected = parsed.get("sections", {})
    order = ["summary", "experience", "education", "skills", "projects",
             "certifications", "awards", "achievements", "roles"]
    titles = {"summary": "Summary", "experience": "Experience",
              "education": "Education", "skills": "Skills",
              "projects": "Projects", "certifications": "Certifications",
              "awards": "Awards",
              "achievements": "Achievements and Certifications",
              "roles": "Roles of Responsibility"}
    # Preserve the resume's OWN header text (Kanish finding: his source says
    # 'Achievements' but a hardcoded canonical label rendered 'Awards').
    # Synonyms exist for MATCHING, never for RELABELING.  Long matches are
    # capped (RULE3 starts-with can swallow trailing words into the match).
    src_titles = parsed.get("section_titles") or {}

    def _label_for(key: str) -> str:
        orig = _sanitize_header_label(src_titles.get(key) or "").strip()
        if orig and len(orig) <= 40:
            # A generic 'Portfolio'/'Works' projects header renders as
            # 'Projects' (see _PROJECT_GENERIC_ALIASES); never the raw tag.
            if key == "projects" and orig.lower() in _PROJECT_GENERIC_ALIASES:
                return "Projects"
            # narrow typo correction ('EDUACTION' -> 'Education'); correctly
            # spelled and differently worded headers pass through untouched
            return canonical_header_spelling(orig)
        return titles[key]

    # Multi-column documents keep LITERAL entry grouping: entry-folding,
    # tag-folding and consolidation all assume single-column reading order
    # and scramble interleaved column lines into false entries/titles
    # (stress_twocol). Literal output is incomplete but never garbled.
    # Uses the STRONG signal - a thin chip rail (flagged by the generic
    # detector) still parses fine with the normal pipeline.
    literal = bool(parsed.get("multicolumn_strong"))

    out_sections = []
    # Lines already claimed by a detected section must never be re-scanned
    # by the skills fallback (prevents phantom Skills sections).
    assigned_ids = set(id(l) for ls in detected.values() for l in ls)
    for key in order:
        sec_lines = detected.get(key)
        if sec_lines is None:
            if key == "skills":
                guessed = _guess_skills(lines, assigned_ids)
                if guessed:
                    out_sections.append({"key": "skills", "title": "Skills",
                                         "type": "skills", "items": guessed})
            continue
        if key == "skills":
            out_sections.append({"key": key, "title": _label_for(key),
                                 "type": "skills",
                                 "items": _flatten_skills(sec_lines)})
        elif key == "summary":
            text = re.sub(r"\s+", " ",
                          " ".join(l["text"].strip() for l in sec_lines))
            out_sections.append({"key": key, "title": _label_for(key),
                                 "type": "paragraph", "text": text})
        else:
            entries = _group_entries(sec_lines, literal=literal)
            if not literal:
                entries = _consolidate_entries(_fold_meta_entries(entries))
            out_sections.append({"key": key, "title": _label_for(key),
                                 "type": "entries", "entries": entries})

    # --- unknown / custom sections: keep them, labelled with their own
    #     original header text, never merge their content into another section.
    detected_titles = parsed.get("section_titles") or {}
    for sid in parsed.get("section_order", []):
        if not sid.startswith("custom:"):
            continue
        sec_lines = detected.get(sid)
        if not sec_lines:
            continue
        label = canonical_header_spelling(
            _sanitize_header_label(detected_titles.get(sid)
                                   or sid.split(":", 1)[1].strip().title()))
        has_bullets = any(l.get("bullet") or _is_bullet_line(l) for l in sec_lines)
        if has_bullets:
            entries = _group_entries(sec_lines, literal=literal)
            if not literal:
                entries = _consolidate_entries(_fold_meta_entries(entries))
            out_sections.append({"key": sid, "title": label, "type": "entries",
                                 "entries": entries})
        else:
            text = re.sub(r"\s+", " ",
                          " ".join(l["text"].strip() for l in sec_lines))
            out_sections.append({"key": sid, "title": label,
                                 "type": "paragraph", "text": text})

    return {"name": name, "headline": headline, "contacts": contacts,
            "sections": out_sections}


def _guess_skills(lines: list, exclude_ids: set | None = None) -> list:
    """Fallback skills detection for resumes with NO skills header.

    Only scans lines that were NOT assigned to any detected section -
    bullets and their wrapped continuations inside other sections must
    never be re-invented as a phantom Skills section (stress_twopage:
    'Maintainer, routezen...' became a bogus Skills block AND duplicated
    the OPEN SOURCE bullet). A comma-separated line is kept WHOLE - it
    renders as one row, never one-item-per-row (stress_headers finding).
    """
    exclude = exclude_ids or set()
    for l in lines:
        if id(l) in exclude:
            continue
        t = l["text"]
        if (t.count(",") >= 3 and not DATE_RE.search(t)
                and len(t) > 20 and not EMAIL_RE.search(t)
                and not _is_bullet(t) and not l.get("bullet")):
            return [t.strip().lstrip("\u2022*- \u00b7").strip()]
    return []


def _split_skills(text: str) -> list:
    parts = re.split(r"[,\u2022|;/]|\s{2,}", text)
    seen, out = set(), []
    for p in parts:
        p = p.strip(" .-")
        if p and 1 < len(p) <= 40 and p.lower() not in seen:
            seen.add(p.lower())
            out.append(p)
    return out[:24]


def _flatten_skills(sec_lines: list) -> list:
    items = []
    for l in sec_lines:
        t = l["text"].strip().lstrip("\u2022*- \u00b7").strip()
        if not t:
            continue
        # Source files sometimes carry doubled colons ("Environments:: Docker");
        # collapse them so the label split in template_engine never renders
        # "Environments: : Docker" (C.6).
        t = re.sub(r"::+", ":", t)
        # Keep category-labeled lines ("Programming Languages: C++, JS, TS")
        # and delimited list lines whole so they render like the reference.
        items.append(t)
    seen, uniq = set(), []
    for i in items:
        k = i.lower()
        if k not in seen:
            seen.add(k)
            uniq.append(i)
    return uniq[:30]


_EMPTY_PARENS_RE = re.compile(r"\s*\(\s*\)")


def _strip_empty_parens(s: str) -> str:
    """Remove orphan empty parenthesis pairs left behind after a date was
    extracted from inside them ('Engineer (2021 - Present)' -> 'Engineer ()'
    -> 'Engineer'). Never touches parens with content."""
    return _EMPTY_PARENS_RE.sub(" ", s or "").strip()


def _clean_title_part(part: str) -> str:
    """Drop junk fragments left over after removing a date from a line.

    "CGPA: 7.82 · 1"  -> "CGPA: 7.82"      (orphan day-of-month removed)
    "Role, Company"   -> "Role, Company"   (untouched)
    """
    part = _strip_empty_parens(part)
    toks = [t.strip(" \u00b7") for t in re.split(r"\u00b7|\||;", part) if t.strip()]
    kept = [t for t in toks
            if len(t) > 3 and not t.isdigit() and re.search(r"[A-Za-z]{2}", t)]
    return " \u00b7 ".join(kept)


def _group_entries(sec_lines: list, literal: bool = False) -> list:
    """Group section lines into entries: {title, meta, date, bullets}.

    Ordered rules (first match wins) so boundaries survive:
      1. bullet line                -> bullet of current entry
      2. bare date line             -> date of current (bullet-less) entry
                                       e.g. education "Sep 2023 - Sep 2027"
      3. dated line                 -> NEW entry (role + date on one line),
                                       unless it merely completes the current
                                       bullet-less one (e.g. "CGPA: 7.82 · 1
                                       Sep 2023 – Sep 2027" after a degree
                                       line) in which case the date attaches.
      4. first title-like line      -> title of fresh entry
      5. title-like line AFTER      -> boundary: flush and start new entry
         content exists                (keeps separate projects separate)
      6. short secondary info       -> meta (CGPA, location), max one
      7. anything else (long prose) -> description bullet

    ``literal=True`` (multi-column documents): title promotion, meta
    attachment and boundary detection are ALL disabled - interleaved
    column lines must never be promoted into false entry titles or
    glued on as meta. The first bold line may open the single entry;
    every other line, bullet or not, becomes a plain bullet.
    """
    entries = []
    cur = {"title": "", "meta": "", "date": "", "bullets": []}
    metas = 0
    title_x = None  # x0 of the current entry's title (for indent bullets)
    # Rail-first date (Muskan): a right-rail date line can precede its own
    # company line while the CURRENT entry already has a date.  Such a date
    # belongs to the NEXT entry - park it here and attach at flush time if
    # that entry has no date of its own.
    pending_date = None

    def flush():
        nonlocal pending_date
        if not cur["date"] and pending_date:
            cur["date"] = pending_date
            pending_date = None
        # a pending date is only consumed when the entry lacks one; if the
        # entry carried its own date the pending one belongs to the NEXT
        # entry (rail-first ordering) and must survive this flush
        if cur["title"] or cur["bullets"]:
            entries.append({
                "title": _strip_empty_parens(cur["title"]),
                "meta": _strip_empty_parens(cur["meta"]),
                "date": cur["date"],
                "bullets": [_clean_bullet(_strip_empty_parens(b))
                            for b in cur["bullets"]],
            })
            if cur.get("link"):
                entries[-1]["link"] = cur["link"]
        cur["title"] = ""
        cur["meta"] = ""
        cur["date"] = ""
        cur["bullets"] = []
        cur.pop("link", None)
        nonlocal metas, title_x
        metas = 0
        title_x = None

    for idx, l in enumerate(sec_lines):
        text = l["text"].strip()
        # consume ANY recognized bullet marker (level-1 and level-2 dots,
        # dashes and Canva private-use variants) - the template adds its own
        # marker glyph, so a leftover leading one would render doubled.
        stripped = strip_bullet_markers(l["text"])
        dm = DATE_RE.search(stripped)
        x0 = l.get("x0_pct")
        is_marker = bool(l.get("bullet") or _is_bullet(text))

        if literal:
            # Multi-column safety net (stress_twocol): lines arrive in
            # interleaved column order, so title/meta/boundary heuristics
            # fabricate nonsense entries. Keep everything literal instead.
            s = stripped.strip()
            if not cur["title"] and l.get("bold") and s and not is_marker:
                cur["title"] = s
            elif s:
                cur["bullets"].append(s)
            continue

        # 0. numbered sub-heading ("1) AlgoZen ...", "2) Codex ...") ----------
        # Project lists are frequently numbered.  These lines can be long
        # (>85 chars with a suffix like an extension ID), so the generic
        # title rules below would misfile them as bullets/prose.  A leading
        # enumerator is always an entry boundary - never a bullet.
        num_heading = (not l.get("bullet")
                       and re.match(r"^\d{1,2}[\)\.\]]\s+\S", stripped)
                       and not dm)
        if num_heading:
            flush()
            cur["title"] = stripped
            title_x = x0
            continue

        # 0.5 short MARKER line acting as a sub-entry title -----------------
        # Design-tool resumes open entries with a level-1 marker bullet
        # ("- Freight Tiger") and put descriptions under deeper level-2
        # markers ("◦ ...") or right-rail tags.  When such a short,
        # unpunctuated level-1 marker line is followed by one of those, it
        # is an ENTRY BOUNDARY - not content of the previous entry.
        lvl1_marker = bool(text) and text.lstrip()[:1] in "\u2022*-"
        if (is_marker and lvl1_marker and not dm and len(stripped) <= 90
                and not stripped.endswith((".", ";"))):
            lvl2_next = False
            for k in range(idx + 1, min(idx + 3, len(sec_lines))):
                t2 = sec_lines[k]["text"].lstrip()
                if t2[:1] in "\u25e6\u25aa\u2023\u25cb\u2218\u25ab\u00b7" \
                        or _rail_fragment_ok(t2) \
                        or sec_lines[k].get("link_uris"):
                    # a hyperlink-carrying line ('GitHub' chip) is a right-rail
                    # element too (Vinod: marker-led project titles demoted to
                    # bullets once GitHub stopped counting as a rail signal)
                    lvl2_next = True
                    break
            if lvl2_next:
                flush()
                cur["title"] = stripped
                uris = l.get("link_uris") or []
                if uris:
                    cur["link"] = uris[0]
                title_x = x0
                print("[entry-promote] %r" % stripped[:40], file=sys.stderr)
                continue

        # 0.9 bare link-label line ("View Project:", "[Report]:", "Link:",
        # and Sarthak's icon-junk chip runs "┬º | \x12 Demo" whose ascii
        # residue is just "Demo") - these carry their target as an overlay
        # hyperlink annotation and must NEVER become titles/bullets/prose.
        # Style/marker/indent do not matter here; Canva renders such chips
        # plain AND bold.
        _chip_probe = _chip_text(stripped)
        if (_LINK_LABEL_RE.fullmatch(stripped)
                or (_LINK_LABEL_RE.fullmatch(_chip_probe)
                    and (l.get("link_uris") or _chip_probe == stripped))):
            uris = l.get("link_uris") or []
            if uris and not cur.get("link"):
                cur["link"] = uris[0]
                print("[entry-link] %r -> %s"
                      % (stripped[:24], uris[0][:60]), file=sys.stderr)
            continue

        # 0.9 bold right-rail tag (Ashmit general fix) -------------------------
        # Dates, metric scores and short tags living in the right rail pair
        # with the CURRENT entry - bold or not.  The old rail path required
        # indent_ok, whose `not bold` clause (a title guard) let BOLD rail
        # fragments ('CGPA: 9.10', '93.7%') fall through to rule 5 and
        # become phantom entry titles.  Position (x0 >= 0.72) plus
        # _rail_fragment_ok is the reliable signal; boldness is not.
        if (not is_marker and x0 is not None and x0 >= 0.72 and cur["title"]
                and _rail_fragment_ok(stripped)):
            if dm and not cur["date"]:
                cur["date"] = stripped.strip("()").strip()
                pending_date = None
            elif dm and cur["date"]:
                # rail-FIRST date for the next entry (Muskan pattern)
                pending_date = stripped.strip("()").strip()
            else:
                cur["meta"] = ((cur["meta"] + " \u00b7 " + stripped)
                               if cur["meta"] else stripped)
            continue

        # 1. bullet --------------------------------------------------------
        # indentation-based list item: clearly right of the entry title, so it
        # cannot be the title / meta / a continuation.  (No title yet -> no
        # indent signal; a section that opens directly with prose still lands
        # in bullets via rule 7.)
        indent_ok = (
            x0 is not None
            and title_x is not None
            and x0 >= title_x + 0.02
            and not l.get("bold")
            and len(stripped) > 1
        )
        if is_marker or indent_ok:
            # A bullet that is just a link label ("Link", "Demo", "Live")
            # carries its target as a hyperlink annotation on the same line -
            # store it as the entry's link instead of a visible bullet.
            if _LINK_LABEL_RE.fullmatch(stripped):
                uris = l.get("link_uris") or []
                if uris:
                    cur["link"] = uris[0]
                continue
            # Right-rail tag on its OWN row ('93%', 'CGPA: 9.07/10', a bare
            # date) pairs with the CURRENT entry - never a description bullet
            # (Vinod: percent scores rendered as bullets, education collapsed)
            if (not is_marker and indent_ok and cur["title"]
                    and x0 >= 0.72 and _rail_fragment_ok(stripped)):
                if dm and not cur["date"]:
                    cur["date"] = stripped.strip("()").strip()
                elif dm and cur["date"]:
                    # date rail-FIRST for the NEXT entry (Muskan: the
                    # '04/06/2025 - 01/07/2025' rail line precedes the
                    # 'BHEL, Haridwar' company line it belongs to)
                    pending_date = stripped.strip("()").strip()
                else:
                    cur["meta"] = ((cur["meta"] + " \u00b7 " + stripped)
                                   if cur["meta"] else stripped)
                continue
            cur["bullets"].append(stripped)
            continue

        date = dm.group(0).strip() if dm else ""
        title_part = ""
        if date:
            title_part = stripped.replace(date, "").strip(" |,\u00b7\u2013\u2014-").strip()

        # 2. bare date line completes the current bullet-less entry --------
        if date and not title_part and cur["title"] \
                and not cur["bullets"] and not cur["date"]:
            # strip wrapping parens: '(2023-Present)' -> '2023-Present'
            cur["date"] = stripped.strip("()").strip()
            pending_date = None
            continue

        # 3. dated line ------------------------------------------------------
        # Fused degree lines ("B.Tech ... CGPA: 8.57 September 2022-Present")
        # run long because the column merger glued the right-rail date on;
        # allow extra room when the text carries a CGPA tag.
        date_len_cap = 135 if "cgpa" in stripped.lower() else 95
        if date and len(stripped) < date_len_cap:
            if cur["title"] and not cur["bullets"]:
                # Continuation of the current entry ("CGPA: 7.82 · <date>"),
                # not the start of a new one - keep them together & clean.
                cur["date"] = date
                pending_date = None
                meta_add = _clean_title_part(title_part)
                if meta_add:
                    cur["meta"] = ((cur["meta"] + " \u00b7 " + meta_add)
                                   if cur["meta"] else meta_add)
                continue
            flush()
            pending_date = None
            cur["title"] = _clean_title_part(title_part)
            cur["date"] = date
            title_x = x0
            continue

        # Parenthetical fragments ("(Working on it)") are annotations, never
        # entry titles - they land in meta (rule 6) / bullets instead.
        looks_title = not (stripped.startswith("(") and stripped.endswith(")")
                           and len(stripped) <= 30)

        # 3.5 role-keyword line (Muskan general fix) ---------------------------
        # Role lines ('Summer Internship-3', 'SDE Intern', 'Backend Developer
        # Intern') are pattern-predictable; company names are not.  A short
        # role line while the entry already carries a NON-role title is that
        # entry's role (meta) - never a new sub-entry boundary, even when
        # meta is already populated (rule 5 would otherwise split the entry
        # and emit the role as a phantom title).
        if (cur["title"] and not cur["bullets"]
                and _ROLE_KW_RE.search(stripped)
                and not _ROLE_KW_RE.search(cur["title"])
                and len(stripped) <= 48
                and not stripped.endswith((".", ";", ":")) and looks_title):
            cur["meta"] = ((cur["meta"] + " \u00b7 " + stripped)
                           if cur["meta"] else stripped)
            metas += 1
            continue

        # 3.6 role-first layout: the entry opened with the ROLE line and the
        # company name follows - swap so the company becomes the title and
        # the role lands in meta (title=company / meta=role convention,
        # matching Kanish 'Freight Tiger' / 'On-Site · SDE Intern').  The
        # length floor keeps short tags like 'Remote' out of the title slot.
        if (cur["title"] and not cur["meta"] and not cur["bullets"]
                and _ROLE_KW_RE.search(cur["title"])
                and not _ROLE_KW_RE.search(stripped)
                and len(stripped) >= 8 and len(stripped) <= 60
                and not stripped.endswith((".", ";", ":")) and looks_title):
            cur["meta"] = cur["title"]
            cur["title"] = stripped
            title_x = x0
            continue

        # 4. first title-like line of a fresh entry ---------------------------
        # The length cap is raised (60 -> 120) for "Name | tech, tech, tech"
        # lines - the pipe title|stack pattern is a strong title signal and
        # Sarthak's 90+ char project titles must not fall through to prose
        # bullets while their tiny chip line takes the title slot.
        if (not cur["title"] and not cur["meta"]
                and (len(stripped) <= 60
                     or ("|" in stripped and len(stripped) <= 120))
                and not stripped.endswith((".", ";", ":")) and looks_title):
            cur["title"] = stripped
            uris = l.get("link_uris") or []
            if uris:
                cur["link"] = uris[0]
                print("[entry-link] %r -> %s"
                      % (stripped[:24], uris[0][:60]), file=sys.stderr)
            title_x = x0
            continue

        # 5. title-like line after content = new sub-entry boundary ----------
        # Same "Name | tech-stack" allowance as rule 4: Sarthak's 90+ char
        # second project title must START A NEW ENTRY, not glue itself onto
        # the previous project as a prose bullet.
        if (len(stripped) <= 85 or ("|" in stripped and len(stripped) <= 120)) \
                and (cur["bullets"] or cur["meta"]) \
                and not stripped.endswith((".", ";")) and looks_title:
            flush()
            cur["title"] = stripped
            uris = l.get("link_uris") or []
            if uris:
                cur["link"] = uris[0]
                print("[entry-link] %r -> %s"
                      % (stripped[:24], uris[0][:60]), file=sys.stderr)
            title_x = x0
            continue

        # 6. secondary info before any bullets (CGPA, location) ---------------
        if not cur["bullets"] and metas < 1 and len(stripped) <= 48 \
                and not stripped.endswith((".", ";", ":")):
            cur["meta"] = (cur["meta"] + " \u00b7 " + stripped) if cur["meta"] \
                else stripped
            metas += 1
            continue

        # 7. long prose -> description ---------------------------------------
        cur["bullets"].append(stripped)

    flush()
    return entries


_TAG_TEXT_RE = re.compile(r"^[A-Z][A-Za-z ,&.'/\\-]{0,47}$")


def _tag_only_entry(e: dict) -> bool:
    """Entry made only of short location/institute/tag fragments ('Punjab,
    India', 'On-Site', 'Thapar Institute of Engineering and Technology')
    - foldable meta material, never a standalone resume entry."""
    if e.get("date") or e.get("link") or e.get("meta"):
        return False
    bl = [b.strip() for b in e.get("bullets", []) if b.strip()]
    t = (e.get("title") or "").strip()
    parts = ([t] if t else []) + bl
    if not parts or len(parts) > 2:
        return False
    if any(len(p) > 48 or not _TAG_TEXT_RE.match(p) for p in parts):
        return False
    return sum(len(p) for p in parts) <= 72


def _tag_frag(e: dict) -> str:
    """Readable "; "-joined fragment of a tag-only entry's texts."""
    bl = [b.strip() for b in e.get("bullets", []) if b.strip()]
    t = (e.get("title") or "").strip()
    return "; ".join([x for x in ([t] if t else []) + bl if x])


def _consolidate_entries(entries: list) -> list:
    """Collapse design-tool column fragments left AFTER grouping.

    Pass 1 (company -> role): a bare short entry ('Freight Tiger', possibly
    holding a right-rail tag like 'On-Site') followed by a dated role entry
    ('SDE Intern' + date + bullets) is one logical job - fold so the render
    shows title=company, meta="Role . Tag", date, role bullets.
    Pass 2 (leading tag): a tag-only entry directly before a real entry
    ("Punjab, India" above the degree row) becomes that entry's meta prefix.
    """
    out = []
    for e in entries:
        prev = out[-1] if out else None
        # ---- pass 1 ------------------------------------------------------
        if (prev is not None and prev.get("title")
                and not prev.get("date") and not prev.get("meta")
                and len(prev["title"]) <= 44
                and all(len(b) <= 30 for b in prev.get("bullets", []))
                and len(prev.get("bullets", [])) <= 2
                and e.get("title") and not e.get("link")):
            parts = [x for x in ([e.get("title", "").strip()]
                                 + [b for b in prev.get("bullets", [])
                                    if b.strip()]) if x]
            if parts:
                e = dict(e)
                # the COMPANY name survives as the entry title; the role line
                # and right-rail tags become the meta ("Role . Tag")
                e["title"] = prev["title"]
                e["meta"] = ((e.get("meta") + " \u00b7 ") if e.get("meta")
                             else "") + " \u00b7 ".join(parts)
                if prev.get("link"):
                    e["link"] = prev["link"]
            print("[entry-fold] %r + %r -> %r / meta %r"
                  % (prev["title"], e.get("title"), e.get("title"),
                     e.get("meta")),
                  file=sys.stderr)
            out.pop()
        # ---- pass 2 ------------------------------------------------------
        elif (out and _tag_only_entry(out[-1])
              and (e.get("date") or e.get("bullets"))):
            t = out[-1]
            frag = _tag_frag(t)
            if frag:
                e = dict(e)
                e["meta"] = ((frag + " \u00b7 " + e["meta"]) if e.get("meta")
                             else frag)
                print("[tag-fold] %r -> meta of %r" % (frag,
                                                       e.get("title")),
                      file=sys.stderr)
                out.pop()
        out.append(e)
    # trailing lone tag-only entries are noise unless they carry bullets
    return [e for e in out if e.get("title") or e.get("bullets")
            or e.get("date")]


def _clean_bullet(b: str) -> str:
    return re.sub(r"\s+", " ", b).strip()


def _fold_meta_entries(entries: list) -> list:
    """Fold short bullet-less trailing entries into the previous entry's meta.

    ``CGPA: 7.82`` parsed as its own (empty) entry is actually secondary info
    about the degree above it.  Folding keeps Education/Experience compact:
    previous -> {title, meta: "CGPA: 7.82", date, bullets}.  Real project
    entries always carry bullets, so they are never folded.
    """
    out = []
    for e in entries:
        title = (e.get("title") or "").strip()
        if (not e.get("bullets") and not e.get("date")
                and not e.get("meta")          # own tag -> its own entry
                and out and out[-1].get("bullets")
                and title and len(title) <= 48):
            prev = out[-1]
            prev["meta"] = ((prev["meta"] + " \u00b7 " + title)
                            if prev.get("meta") else title)
            continue
        out.append(e)
    return out


def structured_to_plain_text(structured: dict) -> str:
    """Flatten structured resume into plain text (for LLM calls / ATS export)."""
    out = [structured.get("name", ""), structured.get("headline", "")]
    out += structured.get("contacts", [])
    for sec in structured.get("sections", []):
        out.append("")
        out.append((sec.get("title") or "").upper())
        if sec.get("type") == "paragraph":
            out.append(sec.get("text", ""))
        elif sec.get("type") == "skills":
            out.append(", ".join(sec.get("items", [])))
        else:
            for e in sec.get("entries", []):
                head = " | ".join(x for x in [e.get("title", ""), e.get("meta", ""),
                                              e.get("date", "")] if x)
                out.append(head)
                out += ["- " + b for b in e.get("bullets", [])]
    return "\n".join(x for x in out if x).strip()


